import logging
import re
from collections.abc import Callable
from pathlib import Path

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)

# 延迟导入 MinerU 客户端，避免在 import 阶段强依赖网络配置
_mineru_client = None


def _get_mineru_client():
    global _mineru_client
    if _mineru_client is None:
        from services.mineru_parser import get_mineru_client

        _mineru_client = get_mineru_client()
    return _mineru_client

MIN_SCAN_TEXT_LENGTH = 50

# 常见简历节标题（用于分节标注）
_SECTION_HEADERS = {
    "个人信息", "基本信息", "个人资料",
    "教育背景", "教育经历", "学历",
    "工作经验", "工作经历", "实习经历", "实习经验",
    "项目经历", "项目经验", "项目",
    "专业技能", "技能", "技术栈",
    "自我评价", "个人评价", "总结",
    "荣誉奖项", "获奖情况", "证书",
    "语言能力", "语言",
    "兴趣爱好", "特长",
    "求职意向", "意向岗位",
}

# 预编译正则：行首匹配节标题（整行完全一致，忽略空白）
_SECTION_RE = re.compile(
    r"^[\s]*(" + "|".join(re.escape(h) for h in _SECTION_HEADERS) + r")[\s]*$",
    re.MULTILINE,
)


def _annotate_sections(text: str) -> str:
    """对文本中的节标题前加 `## ` 标注。

    已标注的行不会重复添加。
    """

    def _replace(match: re.Match) -> str:
        line = match.group(0)
        # 如果已经以 ## 开头，不重复添加
        stripped = line.strip()
        if stripped.startswith("## "):
            return line
        return line.replace(stripped, f"## {stripped}", 1)

    return _SECTION_RE.sub(_replace, text)


# ── 双栏简历阅读顺序重建（Phase 2.5 / P1）──────────────────
# 参考 third_party/SmartResume text_extractor.resort_page_text_with_center_location：
# 按文本行中心 y 坐标分组（同一 y 带归一组）→ 组内按 x 排序 → 按 y 从上到下拼接。
# 修复双栏简历"左栏全读完才读右栏"的串行问题；单栏时行顺序与 pdfplumber 原顺序等价。

# 同一行内 x 间隙超过该值视为分栏断点（页面坐标 pt）
_MULTICOLUMN_MIN_GAP = 60.0
# 全页至少出现 N 行分栏断点才判定双栏（避免噪声误判）
_MULTICOLUMN_MIN_SPLIT = 2
# 词间视觉间隙超过该值才补空格（对齐 pdfplumber extract_words 默认 x_tolerance=3）
_LINE_JOIN_X_TOLERANCE = 3.0


def _extract_link_evidence(
    hyperlinks: list[dict], words: list[dict], existing_text: str
) -> list[str]:
    """Preserve clickable PDF annotations as labelled text evidence.

    PDF text extraction returns the visible label but normally drops the URI
    stored in the annotation.  This is common for labels such as “GitHub 主页”
    and used to create broken resume links.  Only safe web/mail/phone schemes
    are accepted; already-visible values are deduplicated.
    """
    evidence: list[str] = []
    seen = {existing_text}
    for link in hyperlinks or []:
        raw_uri = str(link.get("uri") or "").strip()
        if not raw_uri or len(raw_uri) > 2048:
            continue
        lowered = raw_uri.lower()
        if lowered.startswith("mailto:"):
            value = raw_uri[7:]
        elif lowered.startswith("tel:"):
            value = raw_uri[4:]
        elif lowered.startswith(("http://", "https://")):
            value = raw_uri
        else:
            continue
        canonical_value = value.rstrip("/")
        if not value or any(canonical_value in item for item in seen):
            continue

        try:
            x0, x1 = float(link["x0"]), float(link["x1"])
            top, bottom = float(link["top"]), float(link["bottom"])
            label_words = [
                word
                for word in words
                if x0 <= (float(word["x0"]) + float(word["x1"])) / 2 <= x1
                and top <= (float(word["top"]) + float(word["bottom"])) / 2 <= bottom
            ]
            label_words.sort(key=lambda word: float(word["x0"]))
            label = _join_line_words(label_words).strip()
        except (KeyError, TypeError, ValueError):
            label = ""
        if not label:
            label = "链接"
        line = f"{label}: {value}"
        evidence.append(line)
        seen.add(line)
    return evidence


def _group_by_center_y(texts: list[dict]) -> list[list[dict]]:
    """按文本块中心 y 坐标分组（同一 y 带归一组）。纯函数。

    Args:
        texts: [{text, x0, top, x1, bottom}, ...]，来自 page.extract_words()

    Returns:
        行分组列表；每行内元素附 _center_y/_center_x/_height 临时字段，
        由 _rebuild_reading_order_lines / _is_multicolumn 消费后清理。
        容差 = max(该组平均行高 * 0.5, 10)，与 SmartResume 一致。
    """
    if not texts:
        return []
    for t in texts:
        t["_center_y"] = (t["top"] + t["bottom"]) / 2
        t["_center_x"] = (t["x0"] + t["x1"]) / 2
        t["_height"] = max(t["bottom"] - t["top"], 1.0)
    ordered = sorted(texts, key=lambda x: x["_center_y"])
    lines: list[list[dict]] = []
    current: list[dict] = []
    for t in ordered:
        if not current:
            current = [t]
            continue
        avg_h = sum(x["_height"] for x in current) / len(current)
        tolerance = max(avg_h * 0.5, 10)
        if abs(t["_center_y"] - current[0]["_center_y"]) <= tolerance:
            current.append(t)
        else:
            lines.append(current)
            current = [t]
    if current:
        lines.append(current)
    return lines


def _cleanup_temp_fields(texts: list[dict]) -> None:
    """清理分组时附加的临时坐标字段。"""
    for t in texts:
        for key in ("_center_y", "_center_x", "_height"):
            t.pop(key, None)


def _rebuild_reading_order_lines(texts: list[dict]) -> list[list[dict]]:
    """阅读顺序重建：中心 y 分组 → 组内按中心 x 排序 → 返回行列表。

    Returns:
        list[list[dict]]，每行是已按 x 排序的词块列表（临时字段已清理）。
    """
    lines = _group_by_center_y(texts)
    result: list[list[dict]] = []
    for line in lines:
        line.sort(key=lambda x: x["_center_x"])
        _cleanup_temp_fields(line)
        result.append(line)
    return result


def _join_line_words(words: list[dict]) -> str:
    """按视觉间隙拼接一行内的词：间隙 > 阈值补空格，否则直接相连。

    中文相邻字符间隙小 → 不插空格（"教育背景"保持整行，分节标注可命中）；
    英文词间有间隙 → 补空格（与 pdfplumber extract_text 的空白推断一致）。
    """
    if not words:
        return ""
    parts = [words[0]["text"]]
    for i in range(1, len(words)):
        gap = words[i]["x0"] - words[i - 1]["x1"]
        if gap > _LINE_JOIN_X_TOLERANCE:
            parts.append(" " + words[i]["text"])
        else:
            parts.append(words[i]["text"])
    return "".join(parts)


def _looks_corrupted_pdf_text(text: str) -> bool:
    """Detect font-map/control-character corruption from a PDF extractor."""
    if not text:
        return False
    controls = sum(1 for char in text if ord(char) < 32 and char not in "\n\r\t")
    return controls >= max(3, len(text) // 200)


def _parse_pdf_with_pymupdf(path: str) -> str:
    """Fallback extractor for PDFs whose font map breaks pdfplumber text.

    Import lazily so existing deployments without the optional fallback keep
    the pdfplumber/MinerU/Docling/OCR chain and can report a clear warning.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("PyMuPDF unavailable; cannot repair corrupted PDF text: %s", path)
        return ""

    parts: list[str] = []
    document = fitz.open(path)
    try:
        for page in document:
            page_text = page.get_text("text") or ""
            if page_text.strip():
                parts.append(page_text)
            for link in page.get_links() or []:
                uri = str(link.get("uri") or "").strip()
                if uri.lower().startswith(("http://", "https://", "mailto:", "tel:")):
                    parts.append(f"链接: {uri}")
    finally:
        document.close()
    return "\n".join(parts).strip()


def _is_multicolumn(texts: list[dict]) -> bool:
    """启发式双栏检测：同一 y 带内出现 x 大间隙的列断点 ≥ 阈值次数。

    复用 _group_by_center_y 的行分组；任一行内存在 > _MULTICOLUMN_MIN_GAP 的
    x 间隙断点计一次，全页 ≥ _MULTICOLUMN_MIN_SPLIT 次判双栏。
    """
    if len(texts) < 4:
        return False
    lines = _group_by_center_y(texts)
    split_count = 0
    for line in lines:
        line.sort(key=lambda x: x["_center_x"])
        for i in range(1, len(line)):
            gap = line[i]["x0"] - line[i - 1]["x1"]
            if gap > _MULTICOLUMN_MIN_GAP:
                split_count += 1
                break
    _cleanup_temp_fields(texts)
    return split_count >= _MULTICOLUMN_MIN_SPLIT


def parse_pdf(path: str) -> str:
    """用 pdfplumber 逐页提取 PDF 文本，双栏时重建阅读顺序。

    双栏检测（_is_multicolumn）命中 → 词级坐标重建阅读顺序
    （中心 y 分组 → 组内 x 排序 → 拼接），修复"左栏读完才读右栏"的串行问题；
    单栏 → 保持原 extract_text()，结果与旧行为完全一致。
    """
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            words = page.extract_words()
            if _is_multicolumn(words):
                ordered_lines = _rebuild_reading_order_lines(words)
                if ordered_lines:
                    page_text = "\n".join(
                        _join_line_words(line) for line in ordered_lines
                    )
                    if page_text.strip():
                        link_lines = _extract_link_evidence(
                            getattr(page, "hyperlinks", []), words, page_text
                        )
                        parts.append("\n".join([page_text, *link_lines]))
                continue
            text = page.extract_text()
            if text:
                link_lines = _extract_link_evidence(getattr(page, "hyperlinks", []), words, text)
                parts.append("\n".join([text, *link_lines]))
    text = "\n".join(parts).strip()
    if _looks_corrupted_pdf_text(text):
        repaired = _parse_pdf_with_pymupdf(path)
        if repaired and len(repaired) >= max(MIN_SCAN_TEXT_LENGTH, int(len(text) * 0.5)):
            logger.info(
                "pdfplumber text map corrupted; switched to PyMuPDF: path=%s, old=%d, new=%d",
                path,
                len(text),
                len(repaired),
            )
            return repaired
    return text


def parse_docx(path: str) -> str:
    """逐段 + 逐表格提取 Word 文本。"""
    doc = Document(path)
    parts = []

    # 1. 普通段落
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)

    # 2. 表格（逐行逐单元格拼接）
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    return "\n".join(parts).strip()


def parse_docx_with_sections(path: str) -> str:
    """提取 DOCX 文本并标注节标题。"""
    text = parse_docx(path)
    return _annotate_sections(text)


def parse_txt(path: str) -> str:
    """读取纯文本，UTF-8 → GBK 兜底。"""
    with open(path, "rb") as f:
        raw = f.read()
    for enc in ("utf-8", "gbk", "gb2312"):
        try:
            return raw.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace").strip()


_PARSERS: dict[str, Callable[[str], str]] = {
    ".pdf": parse_pdf,
    ".docx": parse_docx_with_sections,
}


def _is_scanned_pdf(path: str, extracted_text: str) -> bool:
    """Phase 2.5 / P2 门控：pdf_checker 判定扫描件（Producer 指纹 + 文本密度）。

    检测失败/异常 → 返回 False，由既有 len(text) < MIN_SCAN_TEXT_LENGTH 逻辑兜底，
    保证正常文本 PDF 行为不变。
    """
    try:
        from utils.pdf_checker import check_pdf_type

        verdict = check_pdf_type(path, extracted_text=extracted_text)
        if not verdict.detected:
            logger.info("pdf_checker 检测失败，按文本长度兜底: %s", path)
            return False
        if verdict.is_scanned:
            logger.info(
                "pdf_checker 判定扫描件: path=%s, reason=%s",
                path, verdict.reason,
            )
        return verdict.is_scanned
    except Exception as e:
        logger.warning("pdf_checker 异常（按原逻辑兜底）: %s: %s", path, e)
        return False


async def parse_resume(path: str) -> str:
    """根据扩展名自动选解析器。

    优先级：本地快速解析（pdfplumber / python-docx，秒级）→ MinerU → Docling 版面增强 → OCR。

    文本型 PDF/DOCX 的本地提取通常在 1 秒内完成，且 PDF 已包含双栏阅读顺序重建；
    因此不再让所有上传先等待 MinerU 网络轮询。只有本地文本不足或判定为扫描件时，
    才调用 MinerU/Docling/OCR，兼顾正常文件速度与复杂文件质量。
    """
    ext = Path(path).suffix.lower()

    # 1. 本地快速解析（秒级）
    text = ""
    parser = _PARSERS.get(ext)
    if parser is not None:
        try:
            text = parser(path)
        except Exception as e:
            logger.warning("本地解析失败: %s: %s", path, e)
    elif ext == ".txt":
        text = parse_txt(path)
    else:
        raise ValueError(f"不支持的文件格式：{ext}")

    # PDF 也做分节标注（DOCX 在 parse_docx_with_sections 里已做）
    if ext == ".pdf":
        text = _annotate_sections(text or "")

    # 2. 本地结果足够（非扫描件）→ 直接返回，不调用任何外部解析服务
    if text and len(text.strip()) >= MIN_SCAN_TEXT_LENGTH and not _is_scanned_pdf(path, text):
        logger.info("本地快速解析成功: path=%s, len=%d", path, len(text.strip()))
        return text.strip()

    # 3. MinerU 在线精准解析（仅复杂/扫描件，失败自动降级）
    if ext in {".pdf", ".docx"}:
        try:
            client = _get_mineru_client()
            mineru_text = await client.parse_file(path)
            if mineru_text and len(mineru_text) >= MIN_SCAN_TEXT_LENGTH:
                logger.info("MinerU 解析成功: path=%s, len=%d", path, len(mineru_text))
                if ext == ".pdf":
                    mineru_text = _annotate_sections(mineru_text)
                return mineru_text
        except Exception as e:
            logger.warning("MinerU 解析失败，fallback 本地版面/OCR: %s", e)

    # 4. Docling 本地版面解析（本地结果不足/疑似扫描件时的版面感知兜底）
    if ext == ".pdf":
        from utils.docling_parser import parse_pdf_with_docling

        docling_text = await parse_pdf_with_docling(path)
        if docling_text and len(docling_text) >= MIN_SCAN_TEXT_LENGTH:
            logger.info(
                "Docling 版面解析成功: path=%s, len=%d",
                path, len(docling_text),
            )
            return _annotate_sections(docling_text)

    # 5. 扫描件 PDF → RapidOCR 本地逐页 OCR 兜底
    if ext == ".pdf":
        from utils.ocr_parser import ocr_pdf

        ocr_text = await ocr_pdf(path)
        if ocr_text and len(ocr_text) >= MIN_SCAN_TEXT_LENGTH:
            logger.info(
                "RapidOCR 扫描件识别成功: path=%s, len=%d",
                path, len(ocr_text),
            )
            return _annotate_sections(ocr_text)

    logger.warning("解析文本过短: path=%s, len=%d", path, len(text or ""))
    raise ValueError("解析文本过短，可能是扫描件 PDF")
